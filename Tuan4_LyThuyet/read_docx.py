from docx import Document

doc = Document(r'c:\Users\Gigabyte\OneDrive - Industrial University of Ho Chi Minh City\Desktop\clothes_ecomere\Lab_KienTrucThietKePhanMem\Tuan4\Diagram_KTPM.docx')

print("=== NỘI DUNG FILE WORD ===\n")
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Đọc tables nếu có
for table in doc.tables:
    print("\n=== TABLE ===")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(" | ".join(row_text))
